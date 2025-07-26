import sys
import os

# --- Ensure the src/ and project root are in sys.path for reliable imports ---
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
src_dir = os.path.join(project_root, "src")
if project_root not in sys.path:
    sys.path.insert(0, project_root)
if src_dir not in sys.path:
    sys.path.insert(0, src_dir)

import streamlit as st
import asyncio
from ai_integrations import AIService, format_two_paragraphs, remove_variant_label_prefix, remove_dates
from utils import get_logger
from config import NUM_VARIANTS_PER_QUERY, CONCURRENT_AI_CALLS

import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START

import json

logger = get_logger(__name__)

# --- Helper Functions for Output Generation and Download (unchanged) ---

def generate_text_output(all_results_with_variants):
    output_str = ""
    for query_result in all_results_with_variants:
        output_str += f"=== Query ID: {query_result['query_id']} ===\n"
        output_str += f"Query: {query_result['query_text']}\n"
        output_str += f"Client Name: {query_result['client_info'].get('name', 'N/A')}\n"
        if st.session_state.show_debug_outputs or len(query_result['client_info'].get('guidelines', '')) < 200:
             output_str += f"Client Guidelines:\n{query_result['client_info'].get('guidelines', 'N/A')}\n\n"
        else:
             output_str += f"Client Guidelines: (See DOCX/CSV for full details)\n\n"

        for i, variant in enumerate(query_result['variants']):
            output_str += f"--- Variant {i+1} (Angle: {variant['angle']}) ---\n"
            output_str += f"Status: {variant['status']}\n"
            output_str += f"\nFinal Answer:\n{variant['final_answer']}\n\n"
            if st.session_state.show_debug_outputs:
                output_str += f"Research Output:\n{variant['research_output']}\n\n"
                output_str += f"Draft Output:\n{variant['draft']}\n\n"
                output_str += f"Negative Constraints Applied (Previous Final Answers):\n{', '.join(variant['negative_constraints_applied'])}\n\n"
            output_str += "---\n\n"
        output_str += "\n\n"
    return output_str.encode('utf-8')

def generate_csv_output(all_results_with_variants):
    df_data = []
    for query_result in all_results_with_variants:
        for i, variant in enumerate(query_result['variants']):
            row = {
                "Query ID": query_result['query_id'],
                "Query Text": query_result['query_text'],
                "Client Name": query_result['client_info'].get('name', 'N/A'),
                "Client Guidelines": query_result['client_info'].get('guidelines', 'N/A'),
                "Variant Number": i + 1,
                "Angle": variant['angle'],
                "Final Answer": variant['final_answer'],
                "Status": variant['status'],
            }
            if st.session_state.show_debug_outputs:
                row["Research Output (Debug)"] = variant['research_output']
                row["Draft Output (Debug)"] = variant['draft']
                row["Negative Constraints Applied (Previous Final Answers)"] = ' '.join(variant['negative_constraints_applied'])
            df_data.append(row)
    df = pd.DataFrame(df_data)
    return df.to_csv(index=False).encode('utf-8')

def generate_docx_output(all_results_with_variants):
    document = Document()
    document.add_heading('HARO Automation Results', 0)

    if st.session_state.client_info_parsed:
        document.add_heading('Client Information Summary', level=1)
        for q_id, info in st.session_state.client_info_parsed.items():
            document.add_paragraph(f"**Query {q_id} Client:** {info.get('name', 'N/A')}")
            guidelines_para = document.add_paragraph(f"Guidelines: {info.get('guidelines', 'N/A')}")
            guidelines_para.runs[0].font.size = Pt(10)
            document.add_paragraph("")
        document.add_page_break()

    for query_result in all_results_with_variants:
        section = document.add_section(WD_SECTION_START.NEW_PAGE)
        
        document.add_heading(f"Query {query_result['query_id']}: {query_result['query_text'][:100]}...", level=1)
        document.add_paragraph(f"**Client Name:** {query_result['client_info'].get('name', 'N/A')}")
        
        document.add_heading('Original Query & Guidelines (for context)', level=3)
        document.add_paragraph(f"**Original HARO Query:**\n{query_result['query_text']}")
        document.add_paragraph(f"**Client-Specific Guidelines:**\n{query_result['client_info'].get('guidelines', 'N/A')}")
        document.add_paragraph("")

        for i, variant in enumerate(query_result['variants']):
            document.add_heading(f"Variant {i+1} (Angle: {variant['angle']})", level=2)
            document.add_paragraph(f"**Status:** {variant['status']}")

            document.add_heading('Final Answer', level=3)
            for paragraph_text in variant['final_answer'].split('\n\n'):
                if paragraph_text.strip():
                    p = document.add_paragraph(paragraph_text.strip())
                    p.paragraph_format.first_line_indent = Pt(0)

            if st.session_state.show_debug_outputs:
                document.add_heading('Research Output (Debug)', level=3)
                document.add_paragraph(variant['research_output']) 
                document.add_heading('Draft Output (Debug)', level=3)
                for paragraph_text in variant['draft'].split('\n\n'):
                    if paragraph_text.strip():
                        document.add_paragraph(paragraph_text.strip())
                document.add_paragraph(f"**Negative Constraints Applied (Previous Final Answers):** {', '.join(variant['negative_constraints_applied'])}")

            document.add_paragraph("---")

    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()

# --- Asynchronous Processing Function (unchanged) ---

async def run_processing(queries, client_info_map, parameters, status_placeholder):
    ai_service = AIService()
    all_query_results = []
    progress_bar = status_placeholder.progress(0)
    status_text = st.empty()
    tasks = []

    for i, query_data in enumerate(queries):
        query_id = query_data["id"]
        query_text = query_data["text"]
        current_client_info = client_info_map.get(query_id, {})

        tasks.append(
            ai_service.process_query_with_variants(
                query_id,
                query_text,
                current_client_info,
                parameters
            )
        )

    processed_count = 0
    for future in asyncio.as_completed(tasks):
        result = await future
        all_query_results.append(result)
        processed_count += 1
        progress = min(100, (processed_count / len(queries)) * 100)
        progress_bar.progress(int(progress))
        status_text.text(f"Processed {processed_count} of {len(queries)} queries. Generating {NUM_VARIANTS_PER_QUERY} variants each.")

    progress_bar.progress(100)
    status_text.text(f"Processing complete for {len(all_query_results)} queries, each with {NUM_VARIANTS_PER_QUERY} variants.")
    st.success("HARO automation finished!")
    await ai_service.close()
    return all_query_results

# --- Main Streamlit Application (Updated for Login) ---

def main():
    st.set_page_config(page_title="HARO Automation Tool", layout="wide")

    # --- Initialize ALL session state variables at the very top ---
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'results' not in st.session_state:
        st.session_state.results = []
    if 'show_debug_outputs' not in st.session_state:
        st.session_state.show_debug_outputs = False
    if 'client_info_parsed' not in st.session_state:
        st.session_state.client_info_parsed = {}
    if 'general_instructions' not in st.session_state:
        st.session_state.general_instructions = "Ensure answers are concise, impactful, and demonstrate deep industry knowledge."

    # Place this after st.set_page_config and before the main content
    if st.session_state.authenticated:
        _, logout_col = st.columns([0.85, 0.15])
        with logout_col:
            if st.button("Logout"):
                st.session_state.authenticated = False
                st.info("Logged out successfully.")
                st.rerun()    

    DWS_LOGO_URL = "https://media.glassdoor.com/sqll/868966/digital-web-solutions-squarelogo-1579870425403.png"

    # Display the logo and title using columns for alignment
    logo_col, title_col = st.columns([0.1, 0.9])
    with logo_col:
        st.image(DWS_LOGO_URL, width=100)
    with title_col:
        st.title("HARO Answer's Automation Tool (Multi-Variant)")

    # --- AUTHENTICATION SECTION ---
    # If not authenticated, show the login form
    if not st.session_state.authenticated:
        st.subheader("Login to Access HARO Tool")
        
        app_credentials_json_str = st.secrets.get("APP_CREDENTIALS", "{}")
        
        try:
            app_credentials = json.loads(app_credentials_json_str)
        except json.JSONDecodeError:
            st.error("Error loading application credentials. Please ensure 'APP_CREDENTIALS' in st.secrets is valid JSON.")
            app_credentials = {}

        input_col, _ = st.columns([0.5, 0.5])
        with input_col:
            username = st.text_input("Username:")
            password = st.text_input("Password:", type="password")

        login_button_col, _ = st.columns([0.3, 0.7])
        with login_button_col:
            if st.button("Login", type="primary"):
                if username in app_credentials and app_credentials[username] == password:
                    st.session_state.authenticated = True
                    st.success("Login successful!")
                    st.rerun() # Changed from st.experimental_rerun()
                else:
                    st.error("Invalid username or password.")
        
        # IMPORTANT: Stop execution here if not authenticated, so the rest of the app doesn't render
        return 

    # --- LOGOUT BUTTON (Only shown if authenticated) ---
    logout_col, _ = st.columns([0.1, 0.9])
    with logout_col:
        if st.button("Logout"):
            st.session_state.authenticated = False
            st.info("Logged out successfully.")
            st.rerun() # Changed from st.experimental_rerun()

    # --- REST OF THE MAIN APP CONTENT (Only displayed if authenticated) ---
    st.markdown("Enter up to 4 HARO queries and their respective client information. The tool will generate **5 distinct variants** for each query.")

    all_queries_inputs = []
    all_client_info_inputs_raw = []

    st.subheader("Query & Client Information Input")

    for i in range(1, 5):
        st.markdown(f"#### Input for Query {i}")
        cols_q_c = st.columns(2)

        with cols_q_c[0]:
            query_key = f"query_input_{i}"
            query_val = st.text_area(
                f"HARO Query {i}:",
                key=query_key,
                height=150,
                placeholder=f"Paste HARO query {i} here...",
                max_chars=1500
            )
            all_queries_inputs.append(query_val)

        with cols_q_c[1]:
            client_key = f"client_info_input_{i}"
            client_val = st.text_area(
                f"Client Name & Specific Guidelines for Query {i}:",
                key=client_key,
                height=150,
                placeholder=f"E.g.,\nDigital Web Solutions\nInstructions:\nEach answer must be a maximum of 2 paragraphs or 200 words.\nUse a unique perspective that hasn’t been shared before.\nKeep it expert-level yet simple enough for a general audience.\nOnly include one jargon term—avoid overcomplicating the response.\n...",
                max_chars=1500
            )
            all_client_info_inputs_raw.append(client_val)

        st.markdown("---")

    st.session_state.show_debug_outputs = st.checkbox(
        "Show AI Research & Draft (Debug Output) and Constraints",
        value=st.session_state.show_debug_outputs,
        help="If checked, the generated Research and Draft stages, along with negative constraints applied, will be shown in the results and downloads."
    )

    st.sidebar.header("Generation Parameters")
    st.session_state.general_instructions = st.sidebar.text_area(
        "General Guidelines for Tone/Style (Applied to ALL queries/variants):",
        value=st.session_state.general_instructions,
        height=150,
        help="These instructions will be passed to Claude and OpenAI for overall guidance on the output tone and style, *in addition* to client-specific guidelines."
    )

    parameters = {
        "general_instructions": st.session_state.general_instructions
    }

    status_placeholder = st.empty()

    button_col, _ = st.columns([0.3, 0.7])

    with button_col:
        if st.button("Start HARO Automation", type="primary"):
            queries_to_process = []
            client_info_map = {}

            for i in range(4):
                query_text = all_queries_inputs[i].strip()
                client_info_text_raw = all_client_info_inputs_raw[i].strip()

                if query_text:
                    query_id = f"Q{i+1}"
                    queries_to_process.append({"id": query_id, "text": query_text})
                    
                    client_name = ""
                    client_guidelines = ""
                    if client_info_text_raw:
                        lines = client_info_text_raw.split('\n', 1)
                        client_name = lines[0].strip()
                        if len(lines) > 1:
                            client_guidelines = lines[1].strip()

                    client_info_map[query_id] = {
                        "name": client_name if client_name else f"Client {i+1} Default",
                        "guidelines": client_guidelines
                    }

            if not queries_to_process:
                status_placeholder.error("Please enter at least one HARO query in any of the input boxes.")
                return

            st.session_state.client_info_parsed = client_info_map

            status_placeholder.info(f"Starting HARO automation for {len(queries_to_process)} queries, generating {NUM_VARIANTS_PER_QUERY} variants each. This may take a while based on API response times.")

            try:
                st.session_state.results = asyncio.run(
                    run_processing(queries_to_process, client_info_map, parameters, status_placeholder)
                )
            except Exception as e:
                st.error(f"An error occurred during automation: {e}")
                logger.exception("Error in main automation flow.")

    if st.session_state.results:
        st.subheader("Generated HARO Responses")
        for query_result in st.session_state.results:
            client_name_display = query_result['client_info'].get('name', 'N/A')
            query_snippet = query_result['query_text'].split('\n')[0][:70] + "..." if query_result['query_text'] else "..."

            with st.expander(f"Query {query_result['query_id']} (Client: {client_name_display}) - {query_snippet}"):
                with st.expander("Show Full Query & Client Guidelines"):
                    st.markdown(f"**Original HARO Query:**\n```\n{query_result['query_text']}\n```")
                    st.markdown(f"**Client Guidelines:**\n```\n{query_result['client_info'].get('guidelines', 'N/A')}\n```")
                    st.markdown("---")

                for i, variant in enumerate(query_result['variants']):
                    st.markdown(f"#### Variant {i+1} (Angle: {variant['angle']})")
                    st.write(f"**Status:** {variant['status']}")
                    st.markdown(f"**Final Answer:**\n{variant['final_answer']}")
                    if st.session_state.show_debug_outputs:
                        with st.expander(f"Debug Info for Variant {i+1}"):
                            st.markdown(f"**Research Output:**\n```\n{variant['research_output']}\n```")
                            st.markdown(f"**Draft Output:**\n```\n{variant['draft']}\n```")
                            st.markdown(f"**Negative Constraints Applied (Previous Final Answers):** {', '.join(variant['negative_constraints_applied'])}")
                    st.markdown("---")

        st.subheader("Download All Results")
        col_dl1, col_dl2, col_dl3 = st.columns(3)

        with col_dl1:
            st.download_button(
                label="Download as TXT",
                data=generate_text_output(st.session_state.results),
                file_name="haro_responses.txt",
                mime="text/plain",
                help="Download all generated responses as a plain text file."
            )

        with col_dl2:
            st.download_button(
                label="Download as CSV",
                data=generate_csv_output(st.session_state.results),
                file_name="haro_responses.csv",
                mime="text/csv",
                help="Download all generated responses as a CSV file for spreadsheet viewing."
            )
        with col_dl3:
            st.download_button(
                label="Download as DOCX",
                data=generate_docx_output(st.session_state.results),
                file_name="haro_responses.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download all generated responses in a Word Document format."
            )

if __name__ == "__main__":
    main()